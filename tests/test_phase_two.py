import json
from dataclasses import replace
from pathlib import Path
from types import SimpleNamespace

from docx import Document
from ebooklib import epub
import pytest

from contextweaver.adapters import (
    MockTranslationAdapter,
    OpenAICompatibleChatTranslationAdapter,
    OpenAITranslationAdapter,
    TranslationAdapter,
)
from contextweaver.knowledge import propose_knowledge
from contextweaver.markdown import plain_text
from contextweaver.models import ContextPacket, Segment, TranslationRecord
from contextweaver.pipeline import (
    build_context,
    active_translations,
    export_audit_repair_batch,
    import_audit_resolutions,
    import_document,
    import_translation_draft,
    init_project,
    segment_document,
    translate_project,
    validate_project,
)
from contextweaver.reference import import_reference, simplify_reference, simplify_reference_outputs
from contextweaver.storage import read_jsonl, write_jsonl


def _project(tmp_path: Path, source: Path) -> Path:
    root = tmp_path / "project"
    init_project(root, "Phase Two", "en", "zh-CN")
    import_document(root, source)
    return root


def test_structured_markdown_preserves_blocks_and_inline_markup(tmp_path: Path) -> None:
    source = tmp_path / "source.md"
    source.write_text(
        "# Title\n\n- One **bold** item\n- Two items\n\n> A *quoted* line.[^1]\n\n[^1]: Note.\n",
        encoding="utf-8",
    )
    root = _project(tmp_path, source)
    _, segments, _ = segment_document(root)
    assert [item.kind for item in segments] == ["list", "blockquote", "footnote"]
    assert "**bold**" in segments[0].raw
    assert "footnote_ref" in segments[1].format_signature


def test_plain_text_excludes_ordered_list_markers_from_content_checks() -> None:
    assert plain_text("12. **A cited work**") == "A cited work"


def test_docx_import(tmp_path: Path) -> None:
    source = tmp_path / "book.docx"
    document = Document()
    document.add_heading("DOCX Chapter", 1)
    paragraph = document.add_paragraph()
    paragraph.add_run("Important").bold = True
    paragraph.add_run(" text")
    document.save(source)
    root = _project(tmp_path, source)
    doc = json.loads((root / "state" / "source_document.json").read_text())
    assert doc["source_format"] == "docx"
    _, segments, _ = segment_document(root)
    assert segments[0].raw == "**Important** text"


def test_epub_import(tmp_path: Path) -> None:
    source = tmp_path / "book.epub"
    book = epub.EpubBook()
    book.set_identifier("test")
    book.set_title("EPUB Book")
    chapter = epub.EpubHtml(title="One", file_name="one.xhtml")
    chapter.content = "<h1>One</h1><p>Hello EPUB.</p>"
    book.add_item(chapter)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = [chapter]
    epub.write_epub(str(source), book)
    root = _project(tmp_path, source)
    _, segments, _ = segment_document(root)
    assert segments[0].text == "Hello EPUB."
    report = json.loads((root / "state" / "import_report.json").read_text())
    assert report["spine_documents"] == 1


def test_knowledge_proposals_have_evidence_and_preserve_edits(tmp_path: Path) -> None:
    source = tmp_path / "source.md"
    source.write_text("# One\n\nMara met Iven.\n\nMara warned Iven.\n", encoding="utf-8")
    root = _project(tmp_path, source)
    segment_document(root)
    glossary, entities = propose_knowledge(root)
    assert {item.term for item in glossary} == {"Iven", "Mara"}
    assert all(len(item.evidence_segment_ids) == 2 for item in entities)
    path = root / "state" / "glossary.csv"
    path.write_text(path.read_text().replace("Mara,,", "Mara,玛拉,"), encoding="utf-8")
    glossary, _ = propose_knowledge(root)
    assert next(item for item in glossary if item.term == "Mara").preferred_translation == "玛拉"


def test_selective_retranslation_creates_revision_chain(tmp_path: Path) -> None:
    source = tmp_path / "source.md"
    source.write_text("# One\n\nAlpha.\n\nBeta.\n", encoding="utf-8")
    root = _project(tmp_path, source)
    _, segments, _ = segment_document(root)
    translate_project(root, MockTranslationAdapter())
    translate_project(
        root, MockTranslationAdapter(), segment_ids={segments[0].id}, reason="terminology-fix"
    )
    records = read_jsonl(root / "state" / "translations.jsonl", TranslationRecord)
    chain = [item for item in records if item.segment_id == segments[0].id]
    assert [item.revision for item in chain] == [1, 2]
    assert chain[1].supersedes == chain[0].id
    assert active_translations(records)[segments[0].id].reason == "terminology-fix"
    assert validate_project(root) == []


class RetryError(Exception):
    status_code = 429
    response = SimpleNamespace(headers={"retry-after": "0.25"})


class FakeResponses:
    def __init__(self) -> None:
        self.calls = 0

    def create(self, **kwargs: object) -> SimpleNamespace:
        self.calls += 1
        if self.calls == 1:
            raise RetryError("rate limited")
        return SimpleNamespace(output_text=json.dumps({"translations": ["译文"]}))


def test_openai_adapter_retries_rate_limits_without_network() -> None:
    responses = FakeResponses()
    sleeps: list[float] = []
    adapter = OpenAITranslationAdapter(
        client=SimpleNamespace(responses=responses), max_retries=2, sleep=sleeps.append
    )
    segment = Segment(
        "seg", "doc", "sec", 0, "Source", raw="**Source**", format_signature=["strong"]
    )
    packet = ContextPacket("unit", [segment], None, None, None, [], [])
    assert adapter.translate(packet) == ["译文"]
    assert responses.calls == 2
    assert sleeps == [0.25]


def test_compatible_chat_adapter_returns_json_translations() -> None:
    class Completions:
        def create(self, **kwargs: object) -> SimpleNamespace:
            assert kwargs["response_format"] == {"type": "json_object"}
            return SimpleNamespace(
                choices=[SimpleNamespace(message=SimpleNamespace(content='{"translations":["译文"]}'))]
            )

    client = SimpleNamespace(chat=SimpleNamespace(completions=Completions()))
    segment = Segment("seg_1", "doc", "sec", 0, "Source")
    packet = ContextPacket("unit", [segment], None, None, None, [], [])
    adapter = OpenAICompatibleChatTranslationAdapter(
        model="compatible-model", base_url="https://example.invalid", client=client
    )
    assert adapter.translate(packet) == ["译文"]


def test_compatible_chat_adapter_restores_dropped_numeric_reference_links() -> None:
    class Completions:
        def create(self, **kwargs: object) -> SimpleNamespace:
            return SimpleNamespace(
                choices=[
                    SimpleNamespace(
                        message=SimpleNamespace(
                            content='```json\n{"translations":["引文。[66]"]}\n```'
                        )
                    )
                ]
            )

    client = SimpleNamespace(chat=SimpleNamespace(completions=Completions()))
    segment = Segment("seg_1", "doc", "sec", 0, "Citation", raw="Citation.[66](notes.html#fn66)")
    packet = ContextPacket("unit", [segment], None, None, None, [], [])
    adapter = OpenAICompatibleChatTranslationAdapter(
        model="compatible-model", base_url="https://example.invalid", client=client
    )
    assert adapter.translate(packet) == ["引文。[66](notes.html#fn66)"]


def test_compatible_chat_adapter_retries_incomplete_response() -> None:
    class Completions:
        def __init__(self) -> None:
            self.calls = 0

        def create(self, **kwargs: object) -> SimpleNamespace:
            self.calls += 1
            translations = ["第一段"] if self.calls == 1 else ["第一段", "第二段"]
            return SimpleNamespace(
                choices=[
                    SimpleNamespace(message=SimpleNamespace(content=json.dumps({"translations": translations})))
                ]
            )

    completions = Completions()
    client = SimpleNamespace(chat=SimpleNamespace(completions=completions))
    packet = ContextPacket(
        "unit",
        [Segment("seg_1", "doc", "sec", 0, "One"), Segment("seg_2", "doc", "sec", 1, "Two")],
        None,
        None,
        None,
        [],
        [],
    )
    sleeps: list[float] = []
    adapter = OpenAICompatibleChatTranslationAdapter(
        model="compatible-model",
        base_url="https://example.invalid",
        client=client,
        sleep=sleeps.append,
    )
    assert adapter.translate(packet) == ["第一段", "第二段"]
    assert completions.calls == 2
    assert sleeps == [1]


def test_compatible_chat_adapter_retries_lost_markdown_emphasis() -> None:
    class Completions:
        def __init__(self) -> None:
            self.calls = 0

        def create(self, **kwargs: object) -> SimpleNamespace:
            self.calls += 1
            text = "标题" if self.calls == 1 else "*标题*"
            return SimpleNamespace(
                choices=[SimpleNamespace(message=SimpleNamespace(content=json.dumps({"translations": [text]})))]
            )

    completions = Completions()
    client = SimpleNamespace(chat=SimpleNamespace(completions=completions))
    adapter = OpenAICompatibleChatTranslationAdapter(
        model="compatible-model", base_url="https://example.invalid", client=client, sleep=lambda _: None
    )
    packet = ContextPacket("unit", [Segment("seg_1", "doc", "sec", 0, "Title", raw="*Title*")], None, None, None, [], [])
    assert adapter.translate(packet) == ["*标题*"]
    assert completions.calls == 2


class FailSecondUnit(TranslationAdapter):
    name = "intermittent"
    model = "test"

    def __init__(self) -> None:
        self.calls = 0

    def translate(self, packet: ContextPacket) -> list[str]:
        self.calls += 1
        if self.calls == 2:
            raise RuntimeError("temporary outage")
        return [f"done:{item.text}" for item in packet.source_segments]


def test_pipeline_resumes_after_adapter_failure(tmp_path: Path) -> None:
    source = tmp_path / "source.md"
    source.write_text("# One\n\nAlpha.\n\nBeta.\n", encoding="utf-8")
    root = _project(tmp_path, source)
    segment_document(root, unit_size=1)
    import pytest

    with pytest.raises(RuntimeError, match="outage"):
        translate_project(root, FailSecondUnit())
    assert translate_project(root, MockTranslationAdapter()) == (1, 1)
    assert validate_project(root) == []


def test_format_and_terminology_validation(tmp_path: Path) -> None:
    source = tmp_path / "source.md"
    source.write_text("# One\n\nUse **ContextWeaver** now.\n", encoding="utf-8")
    root = _project(tmp_path, source)
    segment_document(root)
    translate_project(root, MockTranslationAdapter())
    glossary = root / "state" / "glossary.csv"
    glossary.write_text(
        "term,preferred_translation,allowed_variants,note,source_segment_id,confidence,evidence_segment_ids,status\nContextWeaver,上下文编织器,,,1,,approved\n",
        encoding="utf-8",
    )
    issues = validate_project(root)
    assert "terminology_mismatch" in {item.kind for item in issues}


def test_image_placeholders_do_not_trigger_repeated_prose_warning(tmp_path: Path) -> None:
    source = tmp_path / "source.md"
    source.write_text("# One\n\n![image](one.jpg)\n\n![image](two.jpg)\n", encoding="utf-8")
    root = _project(tmp_path, source)
    segment_document(root)
    translate_project(root, MockTranslationAdapter())

    issues = validate_project(root)

    assert "repeated_source_inconsistent" not in {item.kind for item in issues}


def test_numeric_anchor_validation_is_blocking(tmp_path: Path) -> None:
    source = tmp_path / "source.md"
    source.write_text("# One\n\nOutput rose 25% in 2024.\n", encoding="utf-8")
    root = _project(tmp_path, source)
    _, segments, _ = segment_document(root)
    translate_project(root, MockTranslationAdapter())
    records = root / "state" / "translations.jsonl"
    records.write_text(records.read_text().replace("25%", "20%"), encoding="utf-8")
    issues = validate_project(root)
    assert "numeric_anchor_mismatch" in {item.kind for item in issues}


def test_audit_repair_batch_is_bounded_and_preserves_current_record(tmp_path: Path) -> None:
    source = tmp_path / "source.md"
    source.write_text("# One\n\nOutput rose 25% in 2024.\n", encoding="utf-8")
    root = _project(tmp_path, source)
    _, segments, _ = segment_document(root)
    draft = tmp_path / "draft.jsonl"
    draft.write_text(
        json.dumps({"segment_id": segments[0].id, "translated_text": "产出在2024年上升20%。"})
        + "\n",
        encoding="utf-8",
    )
    import_translation_draft(root, draft, "codex-agent", "test", "initial")
    package = tmp_path / "repair.jsonl"
    assert export_audit_repair_batch(root, package, max_segments=1) == 1
    row = json.loads(package.read_text(encoding="utf-8"))
    assert row["segment_id"] == segments[0].id
    assert row["current_translation"] == "产出在2024年上升20%。"
    assert row["response_contract"]["required_keys"] == ["segment_id", "translated_text"]


def test_source_backed_audit_resolution_is_bound_to_current_record(tmp_path: Path) -> None:
    source = tmp_path / "source.md"
    source.write_text("# One\n\nNo quantity was stated.\n", encoding="utf-8")
    root = _project(tmp_path, source)
    _, segments, _ = segment_document(root)
    draft = tmp_path / "draft.jsonl"
    draft.write_text(
        json.dumps({"segment_id": segments[0].id, "translated_text": "原文未写数值，但译文写作5月。"})
        + "\n",
        encoding="utf-8",
    )
    import_translation_draft(root, draft, "codex-agent", "test", "initial")
    issue = next(item for item in validate_project(root, numeric_mode="strict") if item.severity == "error")
    record = active_translations(read_jsonl(root / "state" / "translations.jsonl", TranslationRecord))[segments[0].id]
    resolution = tmp_path / "resolution.jsonl"
    resolution.write_text(
        json.dumps(
            {
                "segment_id": segments[0].id,
                "translation_record_id": record.id,
                "issue_id": issue.id,
                "rationale": "The target renders a source-backed calendar context.",
                "evidence": "Source and target were compared by the reviewer.",
            }
        )
        + "\n",
        encoding="utf-8",
    )
    assert import_audit_resolutions(root, resolution, "codex-agent", "test") == 1
    resolved = validate_project(root, numeric_mode="strict")
    assert [(item.severity, item.status) for item in resolved if item.id == issue.id] == [
        ("warning", "resolved")
    ]


def test_balanced_numeric_mode_warns_for_target_only_number(tmp_path: Path) -> None:
    source = tmp_path / "source.md"
    source.write_text("# One\n\nNo quantity was stated.\n", encoding="utf-8")
    root = _project(tmp_path, source)
    _, segments, _ = segment_document(root)
    draft = tmp_path / "draft.jsonl"
    draft.write_text(
        json.dumps(
            {"segment_id": segments[0].id, "translated_text": "原文未说明数量，但这里写了5。"},
            ensure_ascii=False,
        )
        + "\n",
        encoding="utf-8",
    )
    import_translation_draft(root, draft, "codex-agent", "test", "numeric-mode")
    relaxed = validate_project(root)
    balanced = validate_project(root, numeric_mode="balanced")
    strict = validate_project(root, numeric_mode="strict")
    assert not [item for item in relaxed if item.kind == "numeric_anchor_mismatch"]
    assert [item.severity for item in balanced if item.kind == "numeric_anchor_mismatch"] == [
        "warning"
    ]
    assert [item.severity for item in strict if item.kind == "numeric_anchor_mismatch"] == [
        "error"
    ]


def test_numeric_anchors_work_next_to_cjk() -> None:
    from contextweaver.validation import (
        _balanced_numeric_anchors,
        _numbers,
        _numeric_anchors,
    )

    assert _numbers("In 1791, close to 90 percent") == ["1791", "90"]
    assert _numbers("1791年，接近90%") == ["1791", "90"]
    assert _balanced_numeric_anchors(["5"]) == _balanced_numeric_anchors(["ordinal:5"])
    assert _numeric_anchors("seven hundred and fifty people") == ["quantity:750"]
    assert _numeric_anchors("more than three hundred million people") == [
        "quantity:300000000"
    ]
    assert _numeric_anchors("fifty-five million messages") == ["quantity:55000000"]
    assert _numeric_anchors("the first century") == _numeric_anchors("最初一个世纪")
    assert _numeric_anchors("First edition: May 2023") == ["2023", "month:5"]
    assert _numeric_anchors("第一版：2023年5月") == ["2023", "month:5"]
    assert _numeric_anchors("第一版：2023年五月") == ["2023", "month:5"]
    assert _numeric_anchors("Published Jan. 2024") == ["2024", "month:1"]
    assert _numeric_anchors("出版于2024年1月") == ["2024", "month:1"]
    assert _numeric_anchors("It may improve 2023 outcomes") == ["2023"]
    assert _numeric_anchors("during the 1950s, 1960s, and 1970s") == [
        "decade:1950",
        "decade:1960",
        "decade:1970",
    ]
    assert _numeric_anchors("20世纪50、60和70年代") == [
        "decade:1950",
        "decade:1960",
        "decade:1970",
    ]
    assert _numeric_anchors("more than 7.5 million") == _numeric_anchors("750多万")
    assert _numeric_anchors("half a million tons") == _numeric_anchors("50万吨")
    assert _numeric_anchors("World War II") == _numeric_anchors("二战")
    assert _numeric_anchors("between the ages of eight and twelve") == [
        "quantity:12",
        "quantity:8",
    ]
    assert _numeric_anchors("the late 2010s") == _numeric_anchors("2010年代后期")
    assert _balanced_numeric_anchors(
        _numeric_anchors("nineteen sixty-four")
    ) == _balanced_numeric_anchors(_numeric_anchors("1964"))
    assert _numeric_anchors("twenty-five people") == _numeric_anchors("二十五人")
    assert _numeric_anchors("thirty-five years") == _numeric_anchors("三十五年")
    assert _numeric_anchors("during the 19th century") == _numeric_anchors("19世纪")
    assert _numeric_anchors("during the eighteenth century") == _numeric_anchors("十八世纪")
    assert _numeric_anchors("第6、7、8章") == ["chapter:6", "chapter:7", "chapter:8"]
    assert _numeric_anchors("Chapter 3, 73") == ["73", "chapter:3"]
    assert "month:5" not in _numeric_anchors("May (1973)")
    assert _numeric_anchors("between one and five shares") == _numeric_anchors("一至五股")
    assert _numeric_anchors("from twenty-one to thirty-seven years") == _numeric_anchors("二十一至三十七岁")
    assert _numeric_anchors("one billion francs") == _numeric_anchors("十亿法郎")
    assert _numeric_anchors("a thousand years") == _numeric_anchors("千余年")
    assert _numeric_anchors("twelve thousand years") == _numeric_anchors("一万二千年")
    assert _numeric_anchors("两千年乃至七千年") == ["quantity:2000", "quantity:7000"]
    assert _numeric_anchors("at a crossroads") == _numeric_anchors("站在十字路口") == []
    assert _balanced_numeric_anchors(
        _numeric_anchors("multiply two seven-digit numbers")
    ) == _balanced_numeric_anchors(_numeric_anchors("将两个七位数相乘"))
    assert _balanced_numeric_anchors(
        _numeric_anchors("for six and a half years")
    ) == _balanced_numeric_anchors(_numeric_anchors("六年半"))
    assert _numeric_anchors("$20-$30 billion") == _numeric_anchors("200亿至300亿美元")
    assert _numeric_anchors("$4 trillion") == _numeric_anchors("4万亿美元")
    assert _numeric_anchors("during COVID-19") == _numeric_anchors("新冠疫情期间") == []
    assert _numeric_anchors("two and two make four") == _numeric_anchors("二加二等于四")
    assert _numeric_anchors("fourteen hours") == _numeric_anchors("十四个小时")
    assert _numeric_anchors("twenty years") == _numeric_anchors("二十年")
    assert _numeric_anchors("a thousand-year struggle") == _numeric_anchors("千年斗争")
    assert _numeric_anchors("a hundred years") == _numeric_anchors("一百年")
    assert _numeric_anchors("one year later") == _numeric_anchors("一年后")
    assert _numeric_anchors("two people") == _numeric_anchors("两个人")
    assert _numeric_anchors("a weaver") == _numeric_anchors("一名织工") == []
    assert _numeric_anchors("nine hours") == _numeric_anchors("九个小时")
    assert _numeric_anchors("the first decade") == _numeric_anchors("头十年")
    assert _numeric_anchors("thousands of documents") == _numeric_anchors("数千份文件")
    assert _numeric_anchors("seven demands") == _numeric_anchors("七项要求")
    assert _numeric_anchors("several decades") == _numeric_anchors("几十年") == []
    assert _numeric_anchors("three decades") == _numeric_anchors("三十年")
    assert _numeric_anchors("six demands") == _numeric_anchors("六项诉求")
    assert _numeric_anchors("hundreds of years") == _numeric_anchors("数百年") == []
    assert _numeric_anchors("two or three decades") == _numeric_anchors("二三十年")
    assert _numeric_anchors("two or three hours") == _numeric_anchors("两三个钟头")
    assert _balanced_numeric_anchors(_numeric_anchors("a hundred and sixty years")) == _balanced_numeric_anchors(_numeric_anchors("160年"))
    assert _numeric_anchors("three or four years") == _numeric_anchors("三四年")
    assert _numeric_anchors("within twenty minutes") == _numeric_anchors("二十分钟")
    assert _numeric_anchors("three countries") == _numeric_anchors("三国")
    assert _numeric_anchors(".93 liters") == _numeric_anchors("0.93升")
    assert _numeric_anchors("at least seventeen different aunes") == _numeric_anchors(
        "至少十七种不同的 aune"
    )


def test_relaxed_numeric_mode_allows_localized_zero_padded_title(tmp_path: Path) -> None:
    project = tmp_path / "book"
    init_project(project, "Numeric title", "en", "zh-CN")
    source = tmp_path / "source.md"
    source.write_text("# Chapter\n\nCharter 08 was published in 2008.\n", encoding="utf-8")
    import_document(project, source)
    segment_document(project)
    translate_project(project, MockTranslationAdapter())
    records = read_jsonl(project / "state" / "translations.jsonl", TranslationRecord)
    paragraph = next(record for record in records if "Charter 08" in record.translated_text)
    records[records.index(paragraph)] = replace(
        paragraph, translated_text="《零八宪章》于2008年公布。"
    )
    write_jsonl(project / "state" / "translations.jsonl", records)

    relaxed = validate_project(project, numeric_mode="relaxed")
    strict = validate_project(project, numeric_mode="strict")

    assert "numeric_anchor_mismatch" not in {item.kind for item in relaxed}
    assert "numeric_anchor_mismatch" in {item.kind for item in strict}


def test_uppercase_slogans_are_not_acronyms() -> None:
    from contextweaver.validation import _acronyms

    assert _acronyms("THE PUBLIC DOES NOT HAVE TO TAKE WHAT IS DISHED OUT") == []
    assert _acronyms("The slogan was KEEP THEIR HANDS OFF, but IBM disagreed.") == ["IBM"]


def test_calendar_month_naturalization_is_not_an_invented_number(tmp_path: Path) -> None:
    source = tmp_path / "source.md"
    source.write_text("# Copyright\n\nFirst edition: May 2023\n", encoding="utf-8")
    root = _project(tmp_path, source)
    _, segments, _ = segment_document(root)
    draft = tmp_path / "draft.jsonl"
    draft.write_text(
        json.dumps(
            {"segment_id": segments[0].id, "translated_text": "第一版：2023年5月"},
            ensure_ascii=False,
        )
        + "\n",
        encoding="utf-8",
    )
    import_translation_draft(root, draft, "codex-agent", "GPT-5", "date-naturalization")
    issues = validate_project(root)
    assert "numeric_anchor_mismatch" not in {item.kind for item in issues}


@pytest.mark.parametrize(
    ("source", "target"),
    [
        ("in the 1980s", "在20世纪80年代"),
        ("in the 1700s", "在18世纪"),
        ("throughout the 1950s and 1960s", "整个20世纪50年代和60年代"),
        ("during the twentieth century", "在20世纪期间"),
        ("in the third century bce", "公元前3世纪"),
        ("in eighteenth-century Britain", "在18世纪的英国"),
        ("from the sixteenth and seventeenth centuries", "从16世纪和17世纪"),
        ("chapters 5 through 9", "第五章至第九章"),
        ("Chapter 7", "第七章"),
        ("January 5, 1914", "1914年1月5日"),
        ("by May 23", "到5月23日"),
        ("the worker died in June", "工人于6月死亡"),
        ("In October of that year", "同年10月"),
        ("World War II", "第二次世界大战"),
        ("around 1.5 million cars", "约150万辆汽车"),
        ("more than one million options", "一百万多个选项"),
        ("five thousand bonuses", "五千笔奖金"),
        ("for more than a millennium", "延续超过一千年"),
        ("for more than a thousand years", "延续超过一千年"),
        ("two hundred thousand residents", "二十万名居民"),
        ("thirteen thousand sheep", "一万三千只羊"),
        ("fifteen hundred people", "一千五百人"),
        ("about a million people", "约一百万人"),
        ("17–18 million people", "1700万至1800万人"),
        ("over five hundred years", "五百多年"),
        ("ten thousand workers", "一万名工人"),
        ("the 5th Regiment", "第5团"),
    ],
)
def test_semantic_numeric_renderings_share_source_anchors(source: str, target: str) -> None:
    from contextweaver.validation import _numeric_anchors

    assert _numeric_anchors(source) == _numeric_anchors(target)


def test_roman_regnal_numbers_are_not_acronyms() -> None:
    from contextweaver.validation import _acronyms

    assert _acronyms("Henry VIII met Richard II") == []


def test_approved_acronym_translation_preserves_semantic_anchor(tmp_path: Path) -> None:
    source = tmp_path / "source.md"
    source.write_text(
        "# One\n\nA US industrial firm.\n\nIts purposes let us proceed unchanged.\n",
        encoding="utf-8",
    )
    root = _project(tmp_path, source)
    _, segments, _ = segment_document(root)
    glossary = root / "state" / "glossary.csv"
    glossary.write_text(
        "term,preferred_translation,allowed_variants,note,source_segment_id,confidence,evidence_segment_ids,status\n"
        f"US,美国,,Geopolitical abbreviation,{segments[0].id},0.99,{segments[0].id},approved\n",
        encoding="utf-8",
    )
    draft = tmp_path / "draft.jsonl"
    draft.write_text(
        "\n".join(
            json.dumps(item, ensure_ascii=False)
            for item in (
                {"segment_id": segments[0].id, "translated_text": "一家美国工业企业。"},
                {"segment_id": segments[1].id, "translated_text": "其宗旨并未改变。"},
            )
        )
        + "\n",
        encoding="utf-8",
    )
    import_translation_draft(root, draft, "codex-agent", "GPT-5", "acronym-localization")
    issues = validate_project(root)
    assert "acronym_missing" not in {item.kind for item in issues}
    assert "terminology_mismatch" not in {item.kind for item in issues}


def test_link_destination_underscores_are_not_emphasis() -> None:
    from contextweaver.markdown import format_signature

    assert format_signature("See [Chapter 1](009_Chapter_002.xhtml).") == ["link"]


def test_human_reference_alignment_and_mainland_draft(tmp_path: Path) -> None:
    source = tmp_path / "source.md"
    source.write_text(
        "# 1 Control over Technology\n\nDigital technology changes work.\n\nA second paragraph.\n",
        encoding="utf-8",
    )
    root = _project(tmp_path, source)
    _, _, units = segment_document(root, unit_size=1)
    reference = tmp_path / "reference.md"
    reference.write_text(
        "# １ 對科技的掌控\n\n數位科技改變工作。\n\n第二段使用軟體與網路。\n", encoding="utf-8"
    )
    assert import_reference(root, reference, "zh-TW", "Human Translator (reference)") == (1, 2, 1)
    packet = build_context(root, units[0])
    assert packet.reference_texts
    count, output = simplify_reference(root)
    assert count == 2
    simplified = output.read_text(encoding="utf-8")
    assert "数字科技" in simplified
    assert "软件与网络" in simplified
    assert "Human Translator (reference)" in simplified
    assert "not a new translation" in simplified
    packet = build_context(root, units[1])
    assert any("软件与网络" in item for item in packet.reference_texts)
    _, paths = simplify_reference_outputs(root, {"epub"})
    assert paths[0].name == "reference-zh-CN.epub"
    assert epub.read_epub(str(paths[0]), options={"ignore_ncx": True}).spine
