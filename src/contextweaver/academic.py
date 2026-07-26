"""Academic-paper PDF rendering with a deliberately small, inspectable surface."""

from __future__ import annotations

import html
import re
from pathlib import Path

from .models import Section, SectionTitleRecord, Segment, SourceDocument, TranslationRecord
from .pipeline import STATE, active_section_titles, active_translations
from .storage import read_json, read_jsonl


def export_academic_pdf(root: Path, content: str = "translated") -> Path:
    """Write a reflowed A4 paper PDF; never attempt to reproduce source pagination.

    The renderer intentionally keeps figure/table placement near its source
    block.  It is a publication draft, not a claim of journal-template
    fidelity.  A complete active translation is required for translated and
    bilingual output, which prevents a partial run from looking final.
    """
    if content not in {"source", "translated", "bilingual"}:
        raise ValueError("content must be source, translated, or bilingual")
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError as error:
        raise RuntimeError("Install ContextWeaver with the 'pdf' extra to render academic PDFs") from error

    source = SourceDocument(**read_json(root / STATE / "source_document.json"))
    sections = read_jsonl(root / STATE / "sections.jsonl", Section)
    segments = read_jsonl(root / STATE / "segments.jsonl", Segment)
    records = read_jsonl(root / STATE / "translations.jsonl", TranslationRecord)
    active = active_translations(records)
    if content != "source" and set(active) != {item.id for item in segments}:
        raise RuntimeError("Academic PDF requires a complete translation; resume translate first")
    titles = active_section_titles(
        read_jsonl(root / STATE / "section_titles.jsonl", SectionTitleRecord)
    )
    article_title = source.title
    if content != "source" and sections and sections[0].id in titles:
        article_title = titles[sections[0].id].translated_title

    chinese_font, latin_font = _register_academic_pdf_fonts(pdfmetrics, UnicodeCIDFont, TTFont)
    styles = getSampleStyleSheet()
    body = ParagraphStyle("CWAcademicBody", parent=styles["BodyText"], fontName=chinese_font, fontSize=10.5, leading=17, spaceAfter=6)
    source_style = ParagraphStyle("CWAcademicSource", parent=body, textColor=colors.HexColor("#555555"), fontSize=8.5, leading=13)
    title_style = ParagraphStyle("CWAcademicTitle", parent=styles["Title"], fontName=chinese_font, fontSize=18, leading=26, alignment=TA_CENTER, spaceAfter=18)
    credit_style = ParagraphStyle(
        "CWAcademicCredit", parent=body, fontSize=9, leading=13, alignment=TA_CENTER, spaceAfter=12
    )
    heading_style = ParagraphStyle("CWAcademicHeading", parent=styles["Heading2"], fontName=chinese_font, fontSize=13, leading=19, spaceBefore=12, spaceAfter=7)
    caption_style = ParagraphStyle("CWAcademicCaption", parent=body, fontSize=9, leading=14, leftIndent=8, textColor=colors.HexColor("#333333"))
    output = root / "output" / "pdf" / f"{content}.pdf"
    output.parent.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(
        str(output),
        pagesize=A4,
        leftMargin=25 * mm,
        rightMargin=25 * mm,
        topMargin=22 * mm,
        bottomMargin=20 * mm,
        title=article_title,
        author=_translator_credit(active) if content != "source" else "ContextWeaver",
    )
    story = [Paragraph(_escape(article_title), title_style)]
    if content != "source":
        story.append(Paragraph(f"译者：{_escape(_translator_credit(active))}", credit_style))
    by_section: dict[str, list[Segment]] = {}
    for item in segments:
        by_section.setdefault(item.section_id, []).append(item)
    for section in sections:
        section_segments = by_section.get(section.id, [])
        if not section_segments:
            continue
        target_title = titles.get(section.id).translated_title if section.id in titles else section.title
        heading = section.title if content == "source" else target_title
        story.append(Paragraph(_escape(heading), heading_style))
        if content == "bilingual" and target_title != section.title:
            story.append(Paragraph(_escape(section.title), source_style))
        for segment in section_segments:
            target = active[segment.id].translated_text if segment.id in active else segment.raw or segment.text
            if content == "source":
                _append_block(root, story, segment, segment.raw or segment.text, body, caption_style, Table, TableStyle, colors, latin_font)
            elif content == "translated":
                _append_block(root, story, segment, target, body, caption_style, Table, TableStyle, colors, latin_font)
            else:
                _append_block(root, story, segment, segment.raw or segment.text, source_style, caption_style, Table, TableStyle, colors, latin_font)
                _append_block(root, story, segment, target, body, caption_style, Table, TableStyle, colors, latin_font)
                story.append(Spacer(1, 5))
    document.build(story, onFirstPage=_page_number, onLaterPages=_page_number, canvasmaker=lambda *args, **kwargs: _AcademicCanvas(*args, chinese_font=chinese_font, **kwargs))
    return output


def export_academic_docx(root: Path, content: str = "translated", profile: str = "zh-cn-academic") -> Path:
    """Write an editable, reflowed academic DOCX from the active records."""
    if content not in {"source", "translated", "bilingual"}:
        raise ValueError("content must be source, translated, or bilingual")
    if profile not in {"zh-cn-academic", "compact"}:
        raise ValueError("profile must be zh-cn-academic or compact")
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    source = SourceDocument(**read_json(root / STATE / "source_document.json"))
    sections = read_jsonl(root / STATE / "sections.jsonl", Section)
    segments = read_jsonl(root / STATE / "segments.jsonl", Segment)
    active = active_translations(read_jsonl(root / STATE / "translations.jsonl", TranslationRecord))
    if content != "source" and set(active) != {item.id for item in segments}:
        raise RuntimeError("Academic DOCX requires a complete translation; resume translate first")
    titles = active_section_titles(
        read_jsonl(root / STATE / "section_titles.jsonl", SectionTitleRecord)
    )
    article_title = source.title
    if content != "source" and sections and sections[0].id in titles:
        article_title = titles[sections[0].id].translated_title
    document = Document()
    section = document.sections[0]
    # A margin-only setup leaves the default Letter page size intact in Word
    # and LibreOffice.  Persist both dimensions so the editable export really
    # is an A4 academic document, not merely an A4-looking template.
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = section.right_margin = Cm(2.5)
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5 if profile == "zh-cn-academic" else 10)
    normal.paragraph_format.line_spacing = 1.55 if profile == "zh-cn-academic" else 1.35
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.bold = False
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style = document.styles["Title"]
    title.add_run(article_title)
    if content != "source":
        credit = document.add_paragraph(f"译者：{_translator_credit(active)}")
        credit.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in credit.runs:
            run.font.size = Pt(9)
        document.core_properties.author = _translator_credit(active)
    by_section: dict[str, list[Segment]] = {}
    for item in segments:
        by_section.setdefault(item.section_id, []).append(item)
    for item in sections:
        scoped = by_section.get(item.id, [])
        if not scoped:
            continue
        target_title = titles.get(item.id).translated_title if item.id in titles else item.title
        heading = item.title if content == "source" else target_title
        document.add_heading(heading, level=min(item.level, 3))
        if content == "bilingual" and target_title != item.title:
            paragraph = document.add_paragraph(item.title)
            paragraph.runs[0].italic = True
        for segment in scoped:
            target = active[segment.id].translated_text if segment.id in active else segment.raw or segment.text
            if content == "source":
                _append_docx_block(document, root, segment, segment.raw or segment.text, profile)
            elif content == "translated":
                _append_docx_block(document, root, segment, target, profile)
            else:
                _append_docx_block(document, root, segment, segment.raw or segment.text, profile, source=True)
                _append_docx_block(document, root, segment, target, profile)
    output = root / "output" / "doc" / f"{content}.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


def _append_docx_block(
    document: object, root: Path, segment: Segment, value: str, profile: str, source: bool = False
) -> None:
    from docx.shared import Cm, Pt

    image = re.fullmatch(r"!\[[^]]*\]\((assets/[^)]+)\)", value.strip())
    if image:
        path = root / "source" / image.group(1)
        if path.exists():
            paragraph = document.add_paragraph()
            paragraph.alignment = 1
            paragraph.add_run().add_picture(str(path), width=Cm(15))
        return
    if segment.kind == "table":
        rows = _markdown_table(value)
        if rows:
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for row_index, row in enumerate(rows):
                for column_index, value_cell in enumerate(row):
                    table_cell = table.cell(row_index, column_index)
                    table_cell.text = html.unescape(value_cell)
                    if row_index == 0:
                        for run in table_cell.paragraphs[0].runs:
                            run.bold = True
            return
    text = re.sub(r"^>\s*", "", value.strip()) if segment.kind == "blockquote" else value.strip()
    text = text.replace("**", "")
    text = re.sub(r"^```(?:math)?\s*|\s*```$", "", text)
    if not text:
        return
    paragraph = document.add_paragraph(text)
    if profile == "compact":
        paragraph.paragraph_format.space_after = Pt(2)
    if source:
        for run in paragraph.runs:
            run.italic = True
            run.font.size = Pt(9)
    if segment.kind == "blockquote":
        paragraph.paragraph_format.left_indent = Cm(0.4)
        paragraph.paragraph_format.space_after = Pt(4)


def _append_block(
    root: Path,
    story: list[object],
    segment: Segment,
    value: str,
    body: object,
    caption: object,
    table_cls: object,
    table_style: object,
    colors: object,
    latin_font: str,
) -> None:
    from reportlab.lib.units import mm
    from reportlab.platypus import Image, Paragraph, Spacer

    image = re.fullmatch(r"!\[[^]]*\]\((assets/[^)]+)\)", value.strip())
    if image:
        path = root / "source" / image.group(1)
        if path.exists():
            picture = Image(str(path), width=150 * mm, height=100 * mm, kind="proportional")
            picture.hAlign = "CENTER"
            story.extend([picture, Spacer(1, 5)])
        return

    if segment.kind == "table":
        rows = _markdown_table(value, latin_font)
        if rows:
            table = table_cls(rows, repeatRows=1, hAlign="LEFT")
            table.setStyle(
                table_style(
                    [
                        ("FONTNAME", (0, 0), (-1, -1), body.fontName),
                        ("FONTSIZE", (0, 0), (-1, -1), 8),
                        ("LEADING", (0, 0), (-1, -1), 11),
                        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#777777")),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EEEEEE")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            story.extend([table, Spacer(1, 8)])
            return
    style = caption if segment.kind == "blockquote" else body
    text = re.sub(r"^>\s*", "", value.strip()) if segment.kind == "blockquote" else value.strip()
    text = text.replace("**", "")
    text = re.sub(r"^```(?:math)?\s*|\s*```$", "", text)
    if text:
        story.append(Paragraph(_mixed_font_markup(text, latin_font).replace("\n", "<br/>"), style))


def _markdown_table(value: str, latin_font: str | None = None) -> list[list[str]]:
    lines = [line.strip() for line in value.splitlines() if line.strip().startswith("|")]
    if len(lines) < 2:
        return []
    rows = []
    for index, line in enumerate(lines):
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if index == 1 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append([_mixed_font_markup(cell, latin_font) if latin_font else _escape(cell) for cell in cells])
    return rows


def _escape(value: str) -> str:
    return html.escape(value, quote=False)


class _AcademicCanvas:
    """Bind page-number rendering to the same embedded Chinese font."""

    def __new__(cls, *args: object, chinese_font: str, **kwargs: object) -> object:
        from reportlab.pdfgen.canvas import Canvas

        canvas = Canvas(*args, **kwargs)
        canvas._contextweaver_chinese_font = chinese_font
        return canvas


def _page_number(canvas: object, document: object) -> None:
    canvas.saveState()
    canvas.setFont(getattr(canvas, "_contextweaver_chinese_font", "STSong-Light"), 8)
    canvas.drawCentredString(document.pagesize[0] / 2, 12 * 2.83465, str(document.page))
    canvas.restoreState()


def _translator_credit(active: dict[str, TranslationRecord]) -> str:
    """Return a visible, evidence-backed translation credit for target drafts."""
    records = [item for item in active.values() if item.adapter != "bibliography-passthrough"]
    pairs = {(item.adapter, item.model) for item in records}
    if len(pairs) == 1:
        adapter, model = next(iter(pairs))
        label = {"codex-agent": "Codex Agent"}.get(adapter, adapter)
        return f"{label} ({model})"
    return "ContextWeaver translation workflow"


def _register_academic_pdf_fonts(pdfmetrics: object, cid_font: object, tt_font: object) -> tuple[str, str]:
    """Embed Songti SC and Times New Roman when those local fonts are available."""
    songti = Path("/System/Library/Fonts/Supplemental/Songti.ttc")
    times = Path("/System/Library/Fonts/Supplemental/Times New Roman.ttf")
    if songti.exists():
        # In Apple's Songti collection, index 6 is Songti SC Regular.  Earlier
        # indexes select bold faces, which would make ordinary Chinese prose
        # visibly heavier than the requested normal Songti weight.
        pdfmetrics.registerFont(tt_font("CW-Songti", str(songti), subfontIndex=6))
        chinese_font = "CW-Songti"
    else:
        pdfmetrics.registerFont(cid_font("STSong-Light"))
        chinese_font = "STSong-Light"
    if times.exists():
        pdfmetrics.registerFont(tt_font("CW-TimesNewRoman", str(times)))
        latin_font = "CW-TimesNewRoman"
    else:
        latin_font = "Times-Roman"
    return chinese_font, latin_font


def _mixed_font_markup(value: str, latin_font: str | None) -> str:
    """Keep Chinese in the paragraph font while assigning Latin runs Times New Roman."""
    if not latin_font:
        return _escape(value)
    pieces: list[str] = []
    cursor = 0
    for match in re.finditer(r"[A-Za-z0-9][A-Za-z0-9._:/+%#?=&@()\-]*", value):
        pieces.append(_escape(value[cursor : match.start()]))
        pieces.append(f'<font name="{latin_font}">{_escape(match.group())}</font>')
        cursor = match.end()
    pieces.append(_escape(value[cursor:]))
    return "".join(pieces)
